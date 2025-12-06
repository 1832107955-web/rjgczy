import re
import os
import base64
import requests
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='宋体', size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

def fetch_mermaid_image(code):
    try:
        # Encode code to base64
        code_bytes = code.encode('utf-8')
        base64_bytes = base64.urlsafe_b64encode(code_bytes)
        base64_str = base64_bytes.decode('ascii')
        
        url = f"https://mermaid.ink/img/{base64_str}"
        print(f"Fetching image from: {url}")
        response = requests.get(url)
        if response.status_code == 200:
            return io.BytesIO(response.content)
        else:
            print(f"Failed to fetch mermaid image: {response.status_code}")
            return None
    except Exception as e:
        print(f"Error fetching mermaid image: {e}")
        return None

def convert_md_to_docx(md_file, docx_file):
    if os.path.exists(docx_file):
        os.remove(docx_file)
        print(f"Deleted existing {docx_file}")

    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    # Read Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    iterator = iter(lines)
    
    in_code_block = False
    code_content = []
    code_type = ""
    
    for line in iterator:
        line = line.rstrip()
        
        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                if code_type == 'mermaid':
                    print("Found mermaid block, converting to image...")
                    mermaid_code = '\n'.join(code_content)
                    image_stream = fetch_mermaid_image(mermaid_code)
                    if image_stream:
                        print("Image fetched successfully, inserting into document...")
                        document.add_picture(image_stream, width=Inches(6))
                        last_paragraph = document.paragraphs[-1] 
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        print("Image fetch failed, inserting code block instead.")
                        p = document.add_paragraph()
                        runner = p.add_run(mermaid_code)
                        runner.font.name = 'Courier New'
                        runner.font.size = Pt(10)
                else:
                    # Normal code block
                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    runner = p.add_run('\n'.join(code_content))
                    runner.font.name = 'Courier New'
                    runner.font.size = Pt(10)
                
                code_content = []
                in_code_block = False
                code_type = ""
            else:
                # Start of code block
                in_code_block = True
                code_type = line[3:].strip()
            continue
            
        if in_code_block:
            code_content.append(line)
            continue

        # Handle Headers
        if line.startswith('# '):
            # Title - Center it, Big Font
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            set_font(run, '黑体', 22, True) # Title size
            
            # Add some spacing after title for the cover effect
            p_spacing = document.add_paragraph()
            p_spacing.paragraph_format.space_after = Pt(100)

        elif line.startswith('## '):
            # Heading 1
            p = document.add_heading(level=1)
            run = p.add_run(line[3:])
            set_font(run, '黑体', 16, True) # Heading 1 size
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)

        elif line.startswith('### '):
            # Heading 2
            p = document.add_heading(level=2)
            run = p.add_run(line[4:])
            set_font(run, '黑体', 14, True) # Heading 2 size
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        elif line.startswith('#### '):
            # Heading 3
            p = document.add_heading(level=3)
            run = p.add_run(line[5:])
            set_font(run, '黑体', 13, True) # Heading 3 size

        # Handle Lists
        elif line.strip().startswith('* ') or line.strip().startswith('- '):
            p = document.add_paragraph(style='List Bullet')
            content = line.strip()[2:]
            # Bold parsing
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, '宋体', 12, True)
                else:
                    run = p.add_run(part)
                    set_font(run, '宋体', 12)

        # Handle Tables
        elif line.startswith('|'):
            if '---' in line: continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not cells: continue
            
            p = document.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        # Handle Normal Text
        elif line.strip():
            # Check for "Key: Value" pairs on cover page
            if "：" in line and len(line) < 50 and "概构" not in line:
                # Likely cover info
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        set_font(run, '宋体', 14, True)
                    else:
                        run = p.add_run(part)
                        set_font(run, '宋体', 14)
            else:
                p = document.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        set_font(run, '宋体', 12, True)
                    else:
                        run = p.add_run(part)
                        set_font(run, '宋体', 12)
        
        elif not line.strip():
            # Empty line - Preserve it!
            document.add_paragraph()

    document.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

import sys

if __name__ == "__main__":
    output_file = 'design_document.docx'
    if len(sys.argv) > 1:
        output_file = sys.argv[1]
    convert_md_to_docx('design_document.md', output_file)
